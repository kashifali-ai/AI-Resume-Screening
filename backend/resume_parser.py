"""Stage 1: extract plain text from an uploaded resume (PDF, DOCX, TXT).

Improvements over the original: PDF via pdfplumber, DOCX tables are read (not
just paragraphs), and image-only / unreadable files raise a clear error instead
of silently yielding empty text.
"""

import io

import pdfplumber
from docx import Document

from logging_config import get_logger

log = get_logger(__name__)


class UnsupportedFileType(Exception):
    """The uploaded file is not a supported resume format."""


class EmptyResumeError(Exception):
    """No extractable text (e.g. a scanned/image-only PDF needing OCR)."""


def extract_text(filename: str, data: bytes) -> str:
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    if ext == "pdf":
        text = _from_pdf(data)
    elif ext == "docx":
        text = _from_docx(data)
    elif ext in ("txt", "md", "text"):
        text = data.decode("utf-8", errors="ignore")
    else:
        raise UnsupportedFileType(
            f"Unsupported file type '.{ext}'. Upload a PDF, DOCX, or TXT resume."
        )

    text = text.strip()
    if not text:
        raise EmptyResumeError(
            "No text could be extracted. If this is a scanned/image PDF, it needs "
            "OCR — please upload a text-based PDF, DOCX, or TXT."
        )
    log.info("Extracted %d characters from '%s'.", len(text), filename)
    return text


def _from_pdf(data: bytes) -> str:
    pages = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            pages.append(page.extract_text() or "")
    return "\n".join(pages)


def _from_docx(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    # Many resumes put skills in tables — read those too.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)
