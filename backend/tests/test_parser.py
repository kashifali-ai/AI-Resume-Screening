"""Parsing: DOCX tables are read; empty/unsupported inputs raise clearly."""

import io

import pytest
from docx import Document

from resume_parser import EmptyResumeError, UnsupportedFileType, extract_text


def _docx_bytes_with_table() -> bytes:
    doc = Document()
    doc.add_paragraph("Aarav Sharma")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Skills"
    table.rows[0].cells[1].text = "Python, FastAPI"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_docx_table_content_is_extracted():
    text = extract_text("resume.docx", _docx_bytes_with_table())
    assert "Python" in text and "FastAPI" in text  # came from the table


def test_txt_is_extracted():
    assert "hello" in extract_text("r.txt", b"hello world")


def test_unsupported_type_raises():
    with pytest.raises(UnsupportedFileType):
        extract_text("resume.exe", b"\x00\x01")


def test_empty_pdf_like_input_raises():
    with pytest.raises(EmptyResumeError):
        extract_text("blank.txt", b"   \n  ")
